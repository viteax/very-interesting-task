import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class WordClient:
    def __init__(self, doc_path: str):
        self.doc_name = doc_path
        self.doc = Document(docx=doc_path)

    def add_heading2(self, title: str, heading_no: int) -> None:
        self.doc.add_heading(
            f"2.{heading_no}. Решения задач на тему «{title}»",
            level=2,
        )
        self.doc.add_paragraph()

    def add_solution(self, no: int, title: str, descr: str, img_path: str) -> None:
        self.doc.add_paragraph(f"«{title}».")
        self.doc.add_paragraph(f"{descr}")
        self.doc.add_paragraph()

        pic_p = self.doc.add_paragraph()
        pic_p.add_run().add_picture(img_path)
        pic_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        label = self.doc.add_paragraph(f"Рисунок 2.{no} — решение задачи «{title}».")
        label.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.doc.add_paragraph()

    def add_page_break(self):
        self.doc.add_page_break()

    def save(self, doc_name: str) -> None:
        os.makedirs("my_docs", exist_ok=True)
        self.doc.save(f"my_docs/{doc_name}")
